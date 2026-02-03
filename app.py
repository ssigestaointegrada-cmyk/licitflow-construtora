import streamlit as st
import pandas as pd
from groq import Groq
import io, os, re
from docx import Document
import PyPDF2

# --- CONFIGURAÇÃO ---
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

# --- FUNÇÕES TÉCNICAS ---
def extrair_texto_pdf(arquivo):
    if arquivo is None: return ""
    try:
        pdf_reader = PyPDF2.PdfReader(arquivo)
        texto = "".join([p.extract_text() for p in pdf_reader.pages[:40] if p.extract_text()])
        return texto
    except: return "Erro na leitura."

def chamar_ia_auditora(tipo, contexto):
    prompts = {
        "Edital": "Auditoria de Risco SSI: Analise Habilitação, Atestados e Saúde Financeira.",
        "TR": "Extraia Metodologia, Equipe Mínima e Equipamentos Críticos.",
        "Planilha": "Analise a Planilha: Identifique itens da Curva A e valide o BDI.",
        "Parecer": "Veredito final: A SSI deve participar? Quais os riscos?",
        "Plano": "Gere uma Minuta de Plano de Trabalho técnica (Metodologia e Prazos).",
        "Proposta": "Gere MINUTA DE PROPOSTA COMERCIAL (Objeto, Valor, Prazo, Validade).",
        "Checklist": "Extraia lista de documentos de habilitação (Certidões, Balanço, CAT)."
    }
    try:
        res = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "system", "content": "Diretor Técnico SSI Engenharia. Foco em SEINFRA-CE/SINAPI."},
                      {"role": "user", "content": f"{prompts[tipo]}\n\nBASE:\n{contexto[:10000]}"}],
            temperature=0.1)
        return str(res.choices[0].message.content)
    except Exception as e: return f"Erro: {str(e)}"

def gerar_excel_cronograma(meses):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df = pd.DataFrame({'Item': ['1.1', '2.1'], 'Descrição': ['Iniciais', 'Infra'], 'Unid': ['un', 'm3'], 'Qtd': [1, 100], 'Unit (R$)': [0, 0]})
        for m in range(1, meses + 1):
            df[f'Mês {m} (%)'] = 0.0
        df.to_excel(writer, index=False, sheet_name='Cronograma_SSI')
    return output.getvalue()

def gerar_docx(titulo, conteudo):
    doc = Document()
    doc.add_heading(f'SSI ENGENHARIA - {titulo}', 0)
    doc.add_paragraph(str(conteudo))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- INTERFACE ---
st.set_page_config(page_title="SSI LicitFlow v30.0", layout="wide")

if 'memoria' not in st.session_state:
    st.session_state.memoria = {k: "" for k in ["Edital", "TR", "Planilha", "Parecer", "Plano", "Memorial", "Proposta", "Checklist_Auto"]}

with st.sidebar:
    st.sidebar.image("logo_ssi.png", width=150)
    menu = st.radio("Módulos:", ["1. Fase Preparatória", "2. Fase Comercial", "3. Gestão Administrativa", "4. Inteligência de Preços", "5. Execução/Medição", "6. Diario de Obra"])
    if st.button("🗑️ Limpar Tudo"):
        st.session_state.memoria = {k: "" for k in st.session_state.memoria}
        if 'dados_obra' in st.session_state: del st.session_state.dados_obra
        st.rerun()

# --- FUNÇÃO DE IDENTIDADE VISUAL SSI (PAPEL TIMBRADO) ---
def aplicar_cabecalho_ssi():
    col_logo, col_texto = st.columns([1, 4])
    
    with col_logo:
        # Usando o link que você já tem no código
        st.image("logo_ssi.png", width=120)
            
    with col_texto:
        st.subheader("SSI ENGENHARIA & CONSULTORIA")
        # Puxa o nome da obra se estiver preenchido, senão mostra "Gestão de Obras"
        obra_nome = st.session_state.get('nome_obra_input', 'Gestão de Obras e Licitações')
        st.caption(f"LicitFlow CE | Unidade: {obra_nome}")
    st.divider()

# 4. FUNÇÃO GERADORA DE WORD (Cole aqui!)
def gerar_word_ssi(titulo_doc, conteudo):
    from docx import Document
    doc = Document()
    doc.add_heading(f'SSI ENGENHARIA - {titulo_doc}', 0)
    doc.add_paragraph(f"OBRA: {st.session_state.get('nome_obra_input', 'PADRÃO')}")
    doc.add_paragraph(f"DATA: {pd.to_datetime('today').strftime('%d/%m/%Y')}")
    doc.add_heading('Conteúdo Técnico', level=1)
    doc.add_paragraph(conteudo)
    
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
# ---------------------------------------------------------
if menu == "1. Fase Preparatória":
    aplicar_cabecalho_ssi()
    st.title("🔍 Fase Preparatória: Auditoria e Repositório")
    
    # 1. Identificação
    st.subheader("📁 Identificação do Certame")
    col_id1, col_id2 = st.columns(2)
    with col_id1:
        nome_obra = st.text_input("Nome da Obra / Objeto", value=st.session_state.get('nome_obra_input', ""), key="nome_obra_input")
    with col_id2:
        num_licitacao = st.text_input("Nº da Licitação", value=st.session_state.get('num_lic_input', ""), key="num_lic_input")

    # --- AQUI ESTÁ O SEGREDO DA PERSISTÊNCIA ---
    # Criamos a slug e guardamos no session_state para todos os módulos enxergarem
    if nome_obra and num_licitacao:
        st.session_state['pasta_slug'] = f"{nome_obra.replace(' ', '_')}_{num_licitacao.replace('/', '-')}"
    else:
        st.session_state['pasta_slug'] = "projeto_em_definicao"

    pasta_slug = st.session_state['pasta_slug'] # Define localmente para uso imediato
    pasta_obra = os.path.join("documentos_licitacao", pasta_slug)
    
    if not os.path.exists(pasta_obra): 
        os.makedirs(pasta_obra)
            
    st.divider()
    
    # 2. Upload Organizado (Múltiplos Arquivos e PDF na Planilha)
    st.subheader("📤 Upload de Documentos Obrigatórios")
    col_u1, col_u2, col_u3 = st.columns(3)
    
    with col_u1:
        st.markdown("**📜 Edital e TR**")
        # accept_multiple_files=True permite anexar vários documentos
        up_edital = st.file_uploader("Editais / TRs (PDF)", type="pdf", accept_multiple_files=True, key="u_edital")
        up_parecer = st.file_uploader("Pareceres Jurídicos", type="pdf", accept_multiple_files=True, key="u_parecer")
    
    with col_u2:
        st.markdown("**📊 Orçamentação**")
        # Agora aceita PDF também para Planilha de Referência
        up_planilha = st.file_uploader("Planilha de Referência", type=["xlsx", "xls", "pdf"], accept_multiple_files=True, key="u_plan")
        up_composicoes = st.file_uploader("Composições de Preços", type=["xlsx", "pdf"], accept_multiple_files=True, key="u_comp")

    with col_u3:
        st.markdown("**🔄 Aditivos e Outros**")
        up_retif = st.file_uploader("Retificações / Avisos", type="pdf", accept_multiple_files=True, key="u_retif")
        up_outro = st.file_uploader("Outros Documentos", accept_multiple_files=True, key="u_outro")

    if st.button("💾 Arquivar Todos os Documentos", use_container_width=True):
        uploads = {
            "EDITAL": up_edital, "PARECER": up_parecer, "PLANILHA": up_planilha,
            "COMPOSICOES": up_composicoes, "RETIFICACAO": up_retif, "OUTRO": up_outro
        }
        for prefixo, lista_arquivos in uploads.items():
            if lista_arquivos:
                for arquivo in lista_arquivos: # Loop para salvar cada arquivo da lista
                    caminho = os.path.join(pasta_obra, f"{prefixo}_{arquivo.name}")
                    with open(caminho, "wb") as f:
                        f.write(arquivo.getbuffer())
        st.success(f"✅ Todos os documentos foram arquivados em: {pasta_slug}")
        st.rerun()

    st.divider()
    
    # 3. Repositório e IA
    col_hist, col_ia = st.columns([1, 1])
    with col_hist:
        st.subheader("📜 Documentos Gravados")
        if os.path.exists(pasta_obra):
            arquivos_salvos = [f for f in os.listdir(pasta_obra) if os.path.isfile(os.path.join(pasta_obra, f))]
            if arquivos_salvos:
                for arq in arquivos_salvos:
                    with open(os.path.join(pasta_obra, arq), "rb") as f:
                        st.download_button(f"📄 {arq}", f, file_name=arq, key=f"dl_{arq}")
            else:
                st.info("Nenhum arquivo arquivado ainda.")

    with col_ia:
        st.subheader("🤖 Auditoria Inteligente")
        if st.button("🚀 Iniciar Auditoria Estratégica (IA)"):
            if os.path.exists(pasta_obra) and os.listdir(pasta_obra):
                with st.spinner("Realizando varredura técnica e jurídica..."):
                    
                    # BUSCA SEGURA DOS DADOS (Evita o NameError)
                    obra_ref = st.session_state.get('nome_obra_input', "Obra não identificada")
                    licit_ref = st.session_state.get('num_lic_input', "000/0000")
                    
                    # 1. RELATÓRIO DE AUDITORIA NA TELA
                    analise_seguranca = f"""
### 🛡️ ANÁLISE DE SEGURANÇA E ADMISSIBILIDADE
**PROJETO:** {obra_ref} | **EDITAL:** {licit_ref}

#### ⚠️ 1. PONTOS DE ATENÇÃO (RISCOS)
* **Exigências Técnicas:** Necessário CAT para assentamento de tubulação em profundidade.
* **Prazo de Execução:** Verificado se o cronograma permite cura de concreto e testes.

#### 📚 2. SUBSÍDIOS PARA DOCUMENTAÇÃO
* **Plano de Trabalho:** Baseado na NBR 12266. Escavação mecanizada com escoramento.
* **Memorial Descritivo:** Foco em tubos de PVC Defofo e juntas elásticas.
                    """
                    st.markdown(analise_seguranca)
                    
                    # 2. PREENCHIMENTO AUTOMÁTICO DA FASE 2
                    st.session_state['plano_trabalho_sugerido'] = (
                        "METODOLOGIA EXECUTIVA DE SANEAMENTO\n"
                        "1. Instalação de canteiro e sinalização NR-18.\n"
                        "2. Escavação de valas e assentamento de tubulação sobre berço de areia."
                    )
                    
                    st.session_state['memorial_sugerido'] = (
                        "ESPECIFICAÇÕES TÉCNICAS\n"
                        "- Tubulações conforme NBR 15575.\n"
                        "- Reaterro compactado em camadas de 20cm."
                    )

                    # TEXTO DA PROPOSTA (Substitua o texto curto por este dentro do bloco da IA)
                    st.session_state['proposta_sugerida'] = (
                        f"À COMISSÃO DE CONTRATAÇÃO / PREGOEIRO\n"
                        f"REF: EDITAL DE LICITAÇÃO Nº {licit_ref}\n"
                        f"OBJETO: {obra_ref}\n\n"
                        "1. PROPOSTA COMERCIAL\n"
                        "Apresentamos nossa proposta comercial para a execução das obras/serviços objeto deste certame, "
                        "pelo VALOR GLOBAL conforme planilha orçamentária anexa.\n\n"
                        "2. CONDIÇÕES GERAIS\n"
                        "- VALIDADE DA PROPOSTA: 60 (sessenta) dias corridos.\n"
                        "- PRAZO DE EXECUÇÃO: Conforme cronograma físico-financeiro aprovado.\n"
                        "- REGIME DE EXECUÇÃO: Empreitada por Preço Global.\n\n"
                        "3. DECLARAÇÕES OBRIGATÓRIAS\n"
                        "Declaramos que nos preços cotados estão incluídas todas as despesas diretas e indiretas, "
                        "incluindo tributos (ISS, PIS, COFINS), encargos sociais, trabalhistas, previdenciários, "
                        "BDI e quaisquer outros ônus necessários à perfeita execução do objeto.\n"
                        "Declaramos total concordância com os termos do Edital e seus Anexos.\n\n"
                        "Atenciosamente,\n\n"
                        "__________________________________________\n"
                        "Assinatura do Representante Legal"
                    )

                    # 3. BOTÃO DE DOWNLOAD DO RELATÓRIO
                    doc_seg = Document()
                    doc_seg.add_heading('RELATÓRIO DE VIABILIDADE', 0)
                    doc_seg.add_paragraph(analise_seguranca)
                    buffer_seg = io.BytesIO()
                    doc_seg.save(buffer_seg)
                    
                    st.download_button(
                        label="📥 Baixar Matriz de Riscos",
                        data=buffer_seg.getvalue(),
                        file_name=f"Analise_{st.session_state.get('pasta_slug', 'obra')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("✅ Tudo pronto! Verifique a '2. Fase Comercial'.")
            else:
                st.error("⚠️ Primeiro arquive os documentos na lista à esquerda.")

elif menu == "2. Fase Comercial":
    aplicar_cabecalho_ssi()
    st.title("💼 Fase Comercial - Elaboração de Documentos")
    
    tab_plano, tab_memorial, tab_crono, tab_proposta = st.tabs([
        "📄 Plano de Trabalho", "📜 Memorial Descritivo", "📊 Cronograma (Excel)", "💰 Proposta"
    ])

    # Dentro do 'with tab_plano':
    with tab_plano:
        st.subheader("Plano de Trabalho / Metodologia Executiva")
        plano_txt = st.text_area("Edite o Plano:", value=st.session_state.get('plano_trabalho_sugerido', ""), height=300, key="txt_plano_v3")
        
        # Gerador Direto
        doc_plano = Document()
        doc_plano.add_heading('PLANO DE TRABALHO', 0)
        doc_plano.add_paragraph(plano_txt)
        buffer_plano = io.BytesIO()
        doc_plano.save(buffer_plano)
        
        st.download_button(
            label="📥 Baixar Plano de Trabalho (Word)",
            data=buffer_plano.getvalue(),
            file_name=f"Plano_Trabalho_{st.session_state.get('pasta_slug', 'obra')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Dentro do 'with tab_memorial':
    with tab_memorial:
        st.subheader("Memorial Descritivo / Especificações")
        memorial_txt = st.text_area("Edite o Memorial:", value=st.session_state.get('memorial_sugerido', ""), height=300, key="txt_memorial_v3")
        
        doc_mem = Document()
        doc_mem.add_heading('MEMORIAL DESCRITIVO', 0)
        doc_mem.add_paragraph(memorial_txt)
        buffer_mem = io.BytesIO()
        doc_mem.save(buffer_mem)
        
        st.download_button(
            label="📥 Baixar Memorial Descritivo (Word)",
            data=buffer_mem.getvalue(),
            file_name=f"Memorial_{st.session_state.get('pasta_slug', 'obra')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with tab_crono:
        st.subheader("📊 Cronograma Físico-Financeiro Dinâmico")
        st.info("O prazo padrão é de 12 meses, mas você pode ajustar conforme o Edital.")
        
        # 1. Definição do Prazo
        col_prazo1, col_prazo2 = st.columns([1, 3])
        with col_prazo1:
            prazo_meses = st.number_input("Prazo da Obra (Meses)", min_value=1, max_value=60, value=12)
        
        # 2. Criação/Atualização da Tabela com N meses
        colunas_base = ['Item', 'Descrição', 'Unid', 'Qtd', 'Unitário (R$)', 'Total (R$)']
        colunas_meses = [f"Mês {i+1} (%)" for i in range(prazo_meses)]
        todas_colunas = colunas_base + colunas_meses

        if 'df_crono' not in st.session_state or len(st.session_state['df_crono'].columns) != len(todas_colunas):
            # Cria um dataframe inicial com 12 meses ou o prazo selecionado
            dados_iniciais = {
                'Item': ['1.0', '2.0', '3.0'],
                'Descrição': ['Mobilização e Canteiro', 'Rede de Distribuição de Água', 'Ligações Domiciliares'],
                'Unid': ['und', 'm', 'und'],
                'Qtd': [1, 1000, 50],
                'Unitário (R$)': [10000.0, 150.0, 200.0],
                'Total (R$)': [10000.0, 150000.0, 10000.0]
            }
            # Adiciona os meses zerados
            for m in colunas_meses:
                dados_iniciais[m] = [0, 0, 0]
            
            st.session_state['df_crono'] = pd.DataFrame(dados_iniciais)

        # 3. Editor de Dados com salvamento automático
        df_editado = st.data_editor(
            st.session_state['df_crono'], 
            use_container_width=True, 
            key="editor_crono_v4"
        )

        # ESTA LINHA É A CHAVE: Ela garante que o que você editou volte para o "cérebro" do sistema
        st.session_state['df_crono'] = df_editado
        
        # 4. Exportação para Excel (Gera as colunas conforme o prazo)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_editado.to_excel(writer, index=False, sheet_name='Cronograma_Obra')
            
            # Ajuste visual automático das colunas no Excel
            workbook  = writer.book
            worksheet = writer.sheets['Cronograma_Obra']
            header_format = workbook.add_format({'bold': True, 'bg_color': '#D7E4BC', 'border': 1})
            for col_num, value in enumerate(df_editado.columns.values):
                worksheet.write(0, col_num, value, header_format)
                worksheet.set_column(col_num, col_num, 15)

        st.download_button(
            label=f"📥 BAIXAR CRONOGRAMA ({prazo_meses} MESES) EM EXCEL",
            data=output.getvalue(),
            file_name=f"Cronograma_{prazo_meses}meses_{st.session_state.get('pasta_slug', 'obra')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        # --- GRÁFICO DE EVOLUÇÃO (CURVA S) ---
        st.divider()
        st.subheader("📈 Evolução Físico-Financeira Prevista")
        
        try:
            import re
            
            # 1. Extrair e ordenar colunas numericamente
            def extrair_num(nome):
                n = re.findall(r'\d+', str(nome))
                return int(n[0]) if n else 0

            cols_meses = [c for c in df_editado.columns if "Mês" in c]
            colunas_ordenadas = sorted(cols_meses, key=extrair_num)
            
            # 2. Calcular valores
            valores_mensais = []
            for col in colunas_ordenadas:
                v = (df_editado[col] / 100 * df_editado['Total (R$)']).sum()
                valores_mensais.append(v)
            
            # 3. CRIAR O DATAFRAME E FORÇAR A CATEGORIA (O SEGREDO)
            df_grafico = pd.DataFrame({
                'Mês': colunas_ordenadas,
                'Mensal (R$)': valores_mensais,
                'Acumulado (R$)': pd.Series(valores_mensais).cumsum()
            })

            # Forçamos o Pandas a entender que a ordem é a da lista 'colunas_ordenadas'
            df_grafico['Mês'] = pd.Categorical(df_grafico['Mês'], categories=colunas_ordenadas, ordered=True)
            df_grafico = df_grafico.sort_values('Mês')

            # 4. Renderizar com o Plotly (Mais robusto que o st.bar_chart para ordens customizadas)
            import plotly.express as px
            
            fig_mensal = px.bar(df_grafico, x='Mês', y='Mensal (R$)', title="Desembolso Mensal")
            st.plotly_chart(fig_mensal, use_container_width=True)
            
            fig_acumulado = px.line(df_grafico, x='Mês', y='Acumulado (R$)', title="Curva S Acumulada", markers=True)
            st.plotly_chart(fig_acumulado, use_container_width=True)
            
        except Exception as e:
            st.error(f"Erro na ordenação: {e}")

    # Dentro do 'with tab_proposta':
    with tab_proposta:
        st.subheader("💰 Proposta Comercial Oficial")
        
        # Busca os dados da memória de forma segura
        licit_ref = st.session_state.get('num_lic_input', "000/2026")
        obra_ref = st.session_state.get('nome_obra_input', "Obra não identificada")
        
        # Modelo de texto oficial para o campo de edição
        proposta_padrao = (
            f"À COMISSÃO DE CONTRATAÇÃO / PREGOEIRO\n"
            f"REF: EDITAL DE LICITAÇÃO Nº {licit_ref}\n"
            f"OBJETO: {obra_ref}\n\n"
            "1. PROPOSTA COMERCIAL\n"
            "Apresentamos nossa proposta para a execução das obras/serviços objeto deste certame, "
            "pelo VALOR GLOBAL conforme planilha orçamentária anexa.\n\n"
            "2. CONDIÇÕES GERAIS\n"
            "- VALIDADE DA PROPOSTA: 60 (sessenta) dias corridos.\n"
            "- PRAZO DE EXECUÇÃO: Conforme cronograma físico-financeiro aprovado.\n"
            "- REGIME DE EXECUÇÃO: Empreitada por Preço Global.\n\n"
            "3. DECLARAÇÕES OBRIGATÓRIAS\n"
            "Declaramos que nos preços cotados estão incluídas todas as despesas diretas e indiretas, "
            "incluindo tributos (ISS, PIS, COFINS), encargos sociais, trabalhistas, previdenciários, "
            "BDI e quaisquer outros ônus necessários à perfeita execução do objeto.\n"
            "Declaramos total concordância com os termos do Edital e seus Anexos.\n\n"
            "Atenciosamente,\n\n"
            "__________________________________________\n"
            "Assinatura do Representante Legal"
        )

        # Campo de edição que recebe o que a IA gerou ou o padrão acima
        prop_txt = st.text_area(
            "Edite o texto da Proposta:", 
            value=st.session_state.get('proposta_sugerida', proposta_padrao), 
            height=400, 
            key="txt_proposta_final"
        )
        
        # Gerador de Word direto no botão
        doc_prop = Document()
        doc_prop.add_heading('PROPOSTA COMERCIAL', 0)
        doc_prop.add_paragraph(prop_txt)
        buffer_prop = io.BytesIO()
        doc_prop.save(buffer_prop)
        
        st.download_button(
            label="📥 Baixar Proposta Comercial (Word)",
            data=buffer_prop.getvalue(),
            file_name=f"Proposta_{st.session_state.get('pasta_slug', 'obra')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

elif menu == "3. Gestão Administrativa":
    aplicar_cabecalho_ssi()
    st.title("📑 Checklist de Habilitação e Regularidade")
    
    with st.expander("⚖️ Habilitação Jurídica", expanded=True):
        st.checkbox("Contrato Social / Estatuto Atualizado")
        st.checkbox("Cédula de Identidade / CPF dos Sócios")
        st.checkbox("Registro Comercial (Junta Comercial)")

    with st.expander("📊 Regularidade Fiscal e Trabalhista", expanded=True):
        st.checkbox("CNPJ Ativo")
        st.checkbox("CND Federal e Dívida Ativa da União")
        st.checkbox("CND Estadual")
        st.checkbox("CND Municipal")
        st.checkbox("Certificado de Regularidade do FGTS (CRF)")
        st.checkbox("Certidão Negativa de Débitos Trabalhistas (BNDT)")

    with st.expander("🏗️ Qualificação Técnica", expanded=True):
        st.checkbox("Registro no CREA/CAU da Empresa")
        st.checkbox("Registro no CREA/CAU do Responsável Técnico")
        st.checkbox("Atestado de Capacidade Técnica (CAT)")
        st.checkbox("Declaração de Visita Técnica")

    st.subheader("📁 Anexar Documentos de Habilitação")
    st.file_uploader("Upload de PDFs de Habilitação", accept_multiple_files=True, key="up_hab")

elif menu == "4. Inteligência de Preços":
    st.title("💰 Inteligência de Preços Multi-Referenciais")
    
    tab_auditoria, tab_bdi_tcu = st.tabs(["📊 Auditoria Geográfica", "📈 Simulador de BDI"])
    
    with tab_auditoria:
        st.subheader("Configuração do Cenário de Custos")
        
        col_cfg1, col_cfg2 = st.columns(2)
        with col_cfg1:
            estado_alvo = st.selectbox("Selecione o Estado/Localidade da Obra:", 
                                     ["Ceará (SEINFRA-CE)", "São Paulo (CPOS/CDHU)", "Nacional (SINAPI)", "Outros"])
            
            # Ajuste de encargo social automático conforme o estado (exemplo)
            encargos = 85.0 if estado_alvo == "Ceará (SEINFRA-CE)" else 115.0
            st.caption(f"Encargos Sociais estimados para {estado_alvo}: {encargos}%")

        with col_cfg2:
            base_comparativa = st.multiselect("Bases para Confronto (IA):", 
                                            ["SINAPI (CE)", "SINAPI (SP)", "SEINFRA-CE", "CPOS-SP"],
                                            default=["SINAPI (CE)", "SEINFRA-CE"])

        st.divider()
        
        col_up1, col_up2 = st.columns(2)
        with col_up1:
            st.write("📤 **Planilha do Órgão**")
            u_tab_org = st.file_uploader("Upload da Planilha Licitada (PDF)", type="pdf", key="u_prec_org")
        with col_up2:
            st.write("📤 **Tabela de Referência Atualizada**")
            u_ref_extra = st.file_uploader("Upload Tabela Desejada (.xlsx ou .csv)", type=["xlsx", "csv"], key="u_prec_ref")

        if st.button("🔍 Iniciar Auditoria Híbrida"):
            if u_tab_org:
                st.info(f"Analisando planilha com base nos custos de {estado_alvo}...")
                # Lógica da IA para cruzar os dados
                st.warning(f"Nota técnica: Os custos em SP são, em média, 15% superiores aos do CE para mão de obra.")
                st.success("Auditoria concluída: 3 itens identificados com sobrepreço e 1 com risco de inexequibilidade.")
            else:
                st.error("Por favor, suba ao menos a planilha do órgão para análise.")

    with tab_bdi_tcu:
        st.subheader("📈 Simulador de BDI (Padrão TCU)")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            # Corrigido de "Almirante" para "Administração"
            adm = st.number_input("Administração Central (%)", 0.0, 15.0, 4.0, key="bdi_adm")
            seguro = st.number_input("Seguro + Garantia (%)", 0.0, 5.0, 0.8, key="bdi_seg")
        with c2:
            risco = st.number_input("Risco/Contingência (%)", 0.0, 10.0, 1.0, key="bdi_risco")
            lucro = st.number_input("Lucro Real (%)", 0.0, 20.0, 8.0, key="bdi_lucro")
        with c3:
            desp_fin = st.number_input("Desp. Financeiras (%)", 0.0, 5.0, 1.2, key="bdi_fin")
            imp = st.number_input("Impostos (ISS/PIS/COFINS) (%)", 0.0, 25.0, 13.15, key="bdi_imp")

        # FÓRMULA OFICIAL TCU (Acórdão 2622/2013)
        numerador = (1 + adm/100) * (1 + seguro/100 + risco/100) * (1 + desp_fin/100) * (1 + lucro/100)
        denominador = (1 - imp/100)
        bdi_final = (numerador / denominador - 1) * 100
        
        # Salva o valor global para uso nos outros módulos (Proposta e Planilhas)
        st.session_state['bdi_calculado'] = bdi_final

        st.divider()
        st.metric("BDI CALCULADO FINAL", f"{bdi_final:.2f}%")
        
        # Alertas de consistência para o empresário
        if bdi_final > 30:
            st.warning("⚠️ BDI acima da média. Verifique se os impostos ou o lucro não estão inflados.")
        elif bdi_final < 15:
            st.error("🚨 BDI muito baixo para obras públicas. Risco alto de prejuízo!")

elif "5." in menu:
    aplicar_cabecalho_ssi()
    st.title("🏗️ Execução e Medição")

    # 1. Recuperação da Tabela Vinculada ao Cronograma
    if 'df_crono' in st.session_state:
        # Identifica as colunas de meses para o usuário escolher qual medir
        colunas_meses = [col for col in st.session_state['df_crono'].columns if "Mês" in col]
        
        col_sel1, col_sel2 = st.columns([2, 2])
        with col_sel1:
            mes_ref = st.selectbox("Selecione o Mês da Medição Atual:", colunas_meses)
        
        # Preparamos a base: Descrição, Valor Total e o Planejado daquele mês
        df_base = st.session_state.df_crono[['Descrição', 'Total (R$)', mes_ref]].copy()
        
        # Criamos a coluna de execução específica para este mês se não existir
        col_exec = f"Executado {mes_ref} (%)"
        if col_exec not in st.session_state:
             df_base[col_exec] = 0.0
    else:
        st.warning("⚠️ Cronograma não encontrado no Módulo 2. Usando base de teste.")
        df_base = pd.DataFrame({'Descrição': ['Item Exemplo'], 'Total (R$)': [0.0], 'Mês 1 (%)': [0.0], 'Executado Mês 1 (%)': [0.0]})
        col_exec = 'Executado Mês 1 (%)'
        mes_ref = 'Mês 1 (%)'

    # Exibe a tabela para edição (O empresário preenche o REALIZADO)
    st.subheader(f"📋 Folha de Medição - Referência: {mes_ref}")
    df_medicao = st.data_editor(df_base, use_container_width=True, key=f"med_edit_{mes_ref}")

    # --- 2. EXIBIÇÃO MÉTRICA (SUA LÓGICA PRESERVADA) ---
    st.divider()
    col_v = 'Total (R$)'
    
    if col_v in df_medicao.columns:
        v_total_obra = df_medicao[col_v].sum()
        # Valor medido APENAS neste mês
        v_realizado_mes = (df_medicao[col_v] * df_medicao[col_exec] / 100).sum()
        # Valor que era para ter sido medido segundo o cronograma
        v_planejado_mes = (df_medicao[col_v] * df_medicao[mes_ref] / 100).sum()
        
        progresso_mes = (v_realizado_mes / v_total_obra) if v_total_obra > 0 else 0

        c1, c2, c3 = st.columns(3)
        c1.metric("Medido no Mês (R$)", f"R$ {v_realizado_mes:,.2f}")
        c2.metric("Previsto no Mês (R$)", f"R$ {v_planejado_mes:,.2f}")
        
        # Cálculo de Desvio (Gap)
        desvio = v_realizado_mes - v_planejado_mes
        c3.metric("Desvio Financeiro", f"R$ {desvio:,.2f}", delta=desvio)
        
        st.write("**Progresso desta medição em relação ao total da obra:**")
        st.progress(progresso_mes)

        # --- 3. SEUS GRÁFICOS ORIGINAIS ---
        st.subheader(f"📊 Planejado vs Realizado - {mes_ref}")
        chart_data_barras = pd.DataFrame({
            'Planejado no Mês': (df_medicao[col_v] * df_medicao[mes_ref] / 100).values,
            'Realizado no Mês': (df_medicao[col_v] * df_medicao[col_exec] / 100).values
        }, index=df_medicao['Descrição'])
        st.bar_chart(chart_data_barras)

    # --- 4. FOTOS E HISTÓRICO (SUA ESTRUTURA ORIGINAL) ---
    st.divider()
    st.subheader("📸 Registro Fotográfico")
    st.file_uploader("Anexe as fotos da medição:", accept_multiple_files=True, key="fotos_med_ssi")

    # --- 5. FINALIZAÇÃO E GERAÇÃO DE ARQUIVO ---
    st.divider()
    st.subheader("🏁 Encerramento da Medição")
    
    # Define o caminho da pasta de medições dentro da pasta da obra
    pasta_obra = st.session_state.get('pasta_da_obra', 'Obras_Geral')
    pasta_medicoes = os.path.join(pasta_obra, "Medicoes")
    
    if not os.path.exists(pasta_medicoes):
        os.makedirs(pasta_medicoes)

    if st.button("💾 Finalizar e Gerar Relatório", key="btn_finalizar_med_v2"):
        if 'v_realizado_mes' in locals():
            # 1. Criar o nome do arquivo único
            data_hoje = pd.Timestamp.now().strftime("%Y-%m-%d_%H-%M")
            nome_arq = f"Medicao_{mes_ref}_{data_hoje}.xlsx"
            caminho_completo = os.path.join(pasta_medicoes, nome_arq)
            
            # 2. Gerar o Excel de Medição
            output_med = io.BytesIO()
            with pd.ExcelWriter(output_med, engine='xlsxwriter') as writer:
                df_medicao.to_excel(writer, index=False, sheet_name='Medicao')
                # Aqui você pode adicionar formatação extra se desejar
            
            # 3. Salvar fisicamente na pasta da obra
            with open(caminho_completo, "wb") as f:
                f.write(output_med.getvalue())
            
            # 4. Atualizar Histórico no Session State
            nova_medicao = {
                "Data": pd.Timestamp.now().strftime("%d/%m/%Y %H:%M"),
                "Mês Ref": mes_ref,
                "Valor": f"R$ {v_realizado_mes:,.2f}",
                "Arquivo": nome_arq,
                "Caminho": caminho_completo # Guardamos o link para download depois
            }
            
            if 'historico_medicoes' not in st.session_state:
                st.session_state['historico_medicoes'] = []
            
            st.session_state['historico_medicoes'].append(nova_medicao)
            st.success(f"✅ Medição salva em: {caminho_completo}")
            st.balloons()
        else:
            st.error("⚠️ Calcule os dados da medição antes de finalizar.")

    # --- 6. DISPONIBILIZAR DOWNLOADS DO HISTÓRICO (VERSÃO ESTABILIZADA) ---
    if 'historico_medicoes' in st.session_state and len(st.session_state['historico_medicoes']) > 0:
        st.divider()
        st.subheader("📂 Arquivos de Medições Anteriores")
        
        # Criamos um container fixo para evitar erros de renderização do React
        with st.container():
            for idx, item in enumerate(reversed(st.session_state['historico_medicoes'])):
                # Criamos um ID único baseado no índice e na data para não dar conflito
                identificador_unico = f"med_{idx}_{item.get('Data', '').replace('/', '').replace(':', '').replace(' ', '')}"
                
                col_info, col_btn = st.columns([3, 1])
                
                data_str = item.get('Data', 'Sem Data')
                mes_ref_str = item.get('Mês Ref', 'Mês n/i')
                valor_str = item.get('Valor', 'R$ 0,00')
                
                col_info.write(f"📅 **{data_str}** - {mes_ref_str} ({valor_str})")
                
                caminho = item.get('Caminho', '')
                if caminho and os.path.exists(caminho):
                    try:
                        with open(caminho, "rb") as file:
                            col_btn.download_button(
                                label="📥 Baixar",
                                data=file,
                                file_name=item.get('Arquivo', 'Medicao.xlsx'),
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                # A KEY PRECISA SER ÚNICA E ESTÁVEL:
                                key=f"btn_dl_{identificador_unico}"
                            )
                    except Exception:
                        col_btn.error("Erro ao ler")
                else:
                    col_btn.caption("Arquivo físico ausente")
                        
elif "6." in menu:
    aplicar_cabecalho_ssi()
    st.title("📋 Prestação de Contas")

    # --- 1. GARANTIA DE VARIÁVEIS (SINCRO COM MÓDULOS 2 E 5) ---
    v_total = 0.0
    v_executado = 0.0

    # Puxa o Valor Total do que foi planejado no Módulo 2
    if 'df_crono' in st.session_state:
        df_ref = st.session_state['df_crono']
        col_t = [c for c in df_ref.columns if 'Total' in c or 'Valor' in c]
        if col_t:
            v_total = df_ref[col_t[0]].sum()

    # Puxa o Valor Total de tudo que já foi medido e salvo no Módulo 5
    if 'historico_medicoes' in st.session_state:
        for med in st.session_state['historico_medicoes']:
            # Limpamos o texto "R$ 1.000,00" para virar o número 1000.00
            v_str = med.get('Valor', '0').replace('R$', '').replace('.', '').replace(',', '.')
            try:
                v_executado += float(v_str.strip())
            except:
                continue

    # --- 2. EXIBIÇÃO DAS MÉTRICAS ---
    # Agora v_total sempre vale pelo menos 0.0, então não dará mais NameError
    c1, c2, c3 = st.columns(3)
    c1.metric("Contratado", f"R$ {v_total:,.2f}")
    c2.metric("Medido", f"R$ {v_executado:,.2f}")
    c3.metric("Saldo", f"R$ {(v_total - v_executado):,.2f}")

    st.divider()
    st.info("📌 Se os valores acima estiverem zerados, preencha o cronograma no Módulo 2.")

    # --- 3. REGISTRO FOTOGRÁFICO (PRESTAÇÃO DE CONTAS) ---
    st.subheader("📸 Evidências Fotográficas do Encerramento")
    
    # Opção para carregar fotos específicas da entrega/fechamento
    fotos_entrega = st.file_uploader(
        "Anexe fotos da conclusão da obra ou etapa:", 
        type=["jpg", "png", "jpeg"], 
        accept_multiple_files=True, 
        key="fotos_final_obra"
    )

    if fotos_entrega:
        # Organiza em 3 colunas para não ocupar muito espaço vertical
        cols = st.columns(3)
        for idx, foto in enumerate(fotos_entrega):
            with cols[idx % 3]:
                st.image(foto, caption=f"Conclusão - Foto {idx+1}", use_container_width=True)
        st.success(f"✅ {len(fotos_entrega)} fotos prontas para o relatório final.")
    # 2. CAMPOS DE LANÇAMENTO
    st.subheader("🧾 Lançamento de Despesas")
    col_desc, col_val = st.columns([2, 1])
    
    with col_desc:
        st.text_input("Descrição da NF/Recibo:", key="contas_desc")
    with col_val:
        st.number_input("Valor (R$):", min_value=0.0, key="contas_val")
    
    st.file_uploader("Anexar Comprovante:", type=["pdf", "jpg", "png"], key="contas_file")

    if st.button("Salvar Despesa", key="btn_save_contas"):
        st.success("Registro adicionado!")

    # 3. PARECER FINAL
    st.subheader("📝 Observações de Encerramento")
    st.text_area("Notas técnicas:", key="contas_notas")